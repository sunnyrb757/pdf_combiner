#!/usr/bin/env python3
"""
Test Utility: create_test_docx.py
Responsibilities:
- Generate dummy .docx files with sample text.
- Create multiple files to test merging and bookmarks.
"""

from pathlib import Path
try:
    from docx import Document
except ImportError:
    print("Error: python-docx not installed. Run 'pip install python-docx'")
    exit(1)

def create_sample_docx(filename: str, content: str, output_dir: Path):
    output_dir.mkdir(parents=True, exist_ok=True)
    file_path = output_dir / filename
    
    doc = Document()
    doc.add_heading(Path(filename).stem, 0)
    doc.add_paragraph(content)
    doc.save(str(file_path))
    print(f"Created {file_path}")

def main():
    test_dir = Path(".tmp/test_inputs")
    
    samples = [
        ("01_Introduction.docx", "This is the first document. It covers the introduction and background."),
        ("02_Middle_Section.docx", "This is the second document. It contains the core details and technical specs."),
        ("03_Conclusion.docx", "This is the final document. It summarizes the findings and lists next steps.")
    ]
    
    for filename, content in samples:
        create_sample_docx(filename, content, test_dir)
        
    print("\nTest files generated in .tmp/test_inputs/")

if __name__ == "__main__":
    main()
