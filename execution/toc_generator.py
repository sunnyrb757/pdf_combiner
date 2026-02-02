#!/usr/bin/env python3
"""
Module D: TOC Generator (toc_generator.py)
Responsibilities:
- Create a .docx file acting as a Table of Contents.
- Accept a list of (title, page_number) tuples.
- Use python-docx to generate a stylized TOC page.
"""

import logging
from pathlib import Path
from typing import List, Tuple
try:
    from docx import Document
    from docx.shared import Pt
    from docx.enum.text import WD_ALIGN_PARAGRAPH
except ImportError:
    print("Error: python-docx not installed. Run 'pip install python-docx'")
    exit(1)

logger = logging.getLogger(__name__)

def generate_toc_docx(toc_entries: List[Tuple[str, int]], output_path: Path):
    """
    Generates a TOC .docx file.
    
    Args:
        toc_entries: List of (Title, StartPage)
        output_path: Path to save the .docx
    """
    doc = Document()
    
    # Title
    title = doc.add_heading('Table of Contents', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_paragraph() # Spacer
    
    # Add Entries
    for title_text, page_num in toc_entries:
        p = doc.add_paragraph()
        p.add_run(title_text).bold = True
        
        # Simple tab-like spacing (basic version)
        dot_leader = " " + ("." * (60 - len(title_text))) + " "
        p.add_run(dot_leader)
        
        p.add_run(str(page_num))
        
    # Save
    doc.save(str(output_path))
    logger.info(f"Generated TOC document at {output_path}")

if __name__ == "__main__":
    logging.basicConfig(level=logging.INFO)
    test_entries = [("Introduction", 2), ("Middle Section", 3), ("Conclusion", 5)]
    generate_toc_docx(test_entries, Path(".tmp/test_toc.docx"))
